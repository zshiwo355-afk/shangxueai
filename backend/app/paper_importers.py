"""题库批量导入解析与校验：Excel (.xlsx) + Word (.docx) + 导入模板生成。

设计要点：
- 解析永远不写库；只生成预览，由 question_imports_api 落 question_import_jobs.rows_json
- 每行独立校验；错误行不阻塞其他行
- 题型支持中英写法：单选/single/SC、多选/multiple/MC、判断/judge/TF、填空/blank/fill、简答/short
- 多选答案：AB / A,B / A、B / A B 都接受
- 判断题答案：对/错 / 是/否 / T/F / true/false / Y/N
- 填空多空：用 | 分隔
- python-docx 在解析 .docx 时按需 import，避免环境未装时整个模块加载失败
"""
from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from typing import Any

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


# ---------------- 题型标准化 ----------------

QUESTION_TYPE_ALIASES = {
    "single": "single", "单选": "single", "单选题": "single", "sc": "single",
    "multiple": "multiple", "多选": "multiple", "多选题": "multiple", "mc": "multiple", "multi": "multiple",
    "judge": "judge", "判断": "judge", "判断题": "judge", "tf": "judge", "true_false": "judge",
    "blank": "blank", "fill": "blank", "填空": "blank", "填空题": "blank",
    "short": "short_answer", "short_answer": "short_answer", "简答": "short_answer", "简答题": "short_answer",
}

JUDGE_TRUE = {"对", "正确", "是", "t", "true", "y", "yes", "1"}
JUDGE_FALSE = {"错", "错误", "否", "不是", "f", "false", "n", "no", "0"}


def _norm_qtype(raw: str) -> str | None:
    key = (raw or "").strip().lower()
    return QUESTION_TYPE_ALIASES.get(key)


def _norm_judge_answer(raw: str) -> str | None:
    text = (raw or "").strip().lower()
    if text in JUDGE_TRUE:
        return "对"
    if text in JUDGE_FALSE:
        return "错"
    return None


# ---------------- 解析结果数据结构 ----------------


@dataclass
class ParsedRow:
    idx: int
    ok: bool
    data: dict[str, Any] | None = None
    errors: list[str] = field(default_factory=list)
    raw: dict[str, Any] | None = None  # 原始单元格内容，便于前端展示

    def to_dict(self) -> dict[str, Any]:
        return {
            "idx": self.idx,
            "ok": self.ok,
            "data": self.data,
            "errors": self.errors,
            "raw": self.raw,
        }


# ---------------- 单行校验 ----------------


def validate_row(data: dict[str, Any]) -> tuple[bool, list[str], dict[str, Any]]:
    """对一行（已粗解析）做完整校验，返回 (ok, errors, normalized_data)。"""
    errors: list[str] = []

    qtype = _norm_qtype(str(data.get("question_type", "")))
    if not qtype:
        errors.append(f"题型「{data.get('question_type', '')}」无法识别")

    stem = str(data.get("stem", "") or "").strip()
    if not stem:
        errors.append("题干不能为空")

    options_raw = data.get("options") or []
    if isinstance(options_raw, str):
        options = [s.strip() for s in re.split(r"[\n|]+", options_raw) if s.strip()]
    else:
        options = [str(o).strip() for o in options_raw if str(o).strip()]

    correct_raw = data.get("correct_answer", "")
    correct: list[str] = []

    score_raw = data.get("score", 5)
    try:
        score = float(score_raw) if score_raw not in (None, "") else 5.0
    except (TypeError, ValueError):
        errors.append(f"分值「{score_raw}」不是合法数字")
        score = 5.0
    if score <= 0:
        errors.append("分值需大于 0")

    # 按题型校验选项与答案
    if qtype in {"single", "multiple"}:
        if len(options) < 2:
            errors.append("选择题至少需要 2 个选项")
        # 把答案转成字母列表（A/B/C…）；也接受直接写选项原文
        ans_letters = _parse_choice_answer(correct_raw, options)
        if not ans_letters:
            errors.append(f"无法识别正确答案「{correct_raw}」")
        else:
            valid_letters = {chr(ord("A") + i) for i in range(len(options))}
            bad = [a for a in ans_letters if a not in valid_letters]
            if bad:
                errors.append(f"正确答案 {','.join(bad)} 不在选项范围内")
            if qtype == "single" and len(ans_letters) != 1 and not bad:
                errors.append("单选题正确答案应只有 1 个")
            correct = ans_letters
    elif qtype == "judge":
        ans = _norm_judge_answer(str(correct_raw))
        if not ans:
            errors.append(f"判断题答案应为 对/错（收到「{correct_raw}」）")
        if not options:
            options = ["对", "错"]
        if ans:
            correct = [ans]
    elif qtype == "blank":
        # 填空：correct_raw 用 | 分多个空，每空可有多个备选答案（用 / 分）
        text = str(correct_raw or "").strip()
        if not text:
            errors.append("填空题需要参考答案")
        else:
            correct = [s.strip() for s in re.split(r"[|\n]+", text) if s.strip()]
        options = []
    elif qtype == "short_answer":
        # 简答：参考答案可空（人工评分）；简答题没有选项
        text = str(correct_raw or "").strip()
        correct = [text] if text else []
        options = []

    normalized = {
        "question_type": qtype or "",
        "stem": stem,
        "options": options,
        "correct_answer": correct,
        "default_score": score,
        "category": str(data.get("category", "") or "").strip(),
        "tag": str(data.get("tag", "") or "").strip(),
        "difficulty": str(data.get("difficulty", "") or "").strip(),
        "explanation": str(data.get("explanation", "") or "").strip(),
    }
    return (not errors), errors, normalized


_CHOICE_LETTER_RE = re.compile(r"^[A-Za-z]$")


def _parse_choice_answer(raw: Any, options: list[str]) -> list[str]:
    """把单/多选答案解析成大写字母列表。

    支持：'A' / 'AB' / 'A,B' / 'A、B' / 'A B' / 中文/英文逗号 / 也接受答案原文。
    """
    if raw is None:
        return []
    if isinstance(raw, (list, tuple, set)):
        items = [str(r).strip() for r in raw if str(r).strip()]
    else:
        text = str(raw).strip()
        if not text:
            return []
        # 拆分
        parts = re.split(r"[\s,，、|/]+", text)
        if len(parts) == 1 and len(text) > 1 and all(_CHOICE_LETTER_RE.match(c) for c in text):
            parts = list(text)
        items = [p.strip() for p in parts if p.strip()]

    letters: list[str] = []
    for it in items:
        if _CHOICE_LETTER_RE.match(it):
            letters.append(it.upper())
            continue
        # 答案是选项原文
        for i, opt in enumerate(options):
            if opt and (opt == it or opt.strip() == it.strip()):
                letters.append(chr(ord("A") + i))
                break
    # 去重保持顺序
    seen = set()
    out: list[str] = []
    for l in letters:
        if l not in seen:
            out.append(l)
            seen.add(l)
    return out


# ---------------- Excel 解析 ----------------


_EXCEL_HEADERS = [
    "题型*", "题干*", "选项A", "选项B", "选项C", "选项D", "选项E", "选项F",
    "正确答案*", "分值*", "分类", "标签", "难度", "解析",
]


def parse_excel(file_bytes: bytes) -> list[ParsedRow]:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    # 跳过表头
    data_rows = rows[1:]
    out: list[ParsedRow] = []
    for i, row in enumerate(data_rows):
        # 跳过整行空行
        if not row or all(c is None or (isinstance(c, str) and not c.strip()) for c in row):
            continue

        cells = list(row) + [None] * max(0, 14 - len(row))
        qtype, stem, oa, ob, oc, od, oe, of, correct, score, category, tag, difficulty, explanation = cells[:14]

        options = [o for o in [oa, ob, oc, od, oe, of] if o not in (None, "")]
        raw = {
            "question_type": qtype,
            "stem": stem,
            "options": [str(o).strip() if o is not None else "" for o in [oa, ob, oc, od, oe, of]],
            "correct_answer": correct,
            "score": score,
            "category": category,
            "tag": tag,
            "difficulty": difficulty,
            "explanation": explanation,
        }
        ok, errors, normalized = validate_row(raw)
        out.append(ParsedRow(idx=i, ok=ok, data=normalized, errors=errors, raw=raw))
    return out


# ---------------- Word 解析（轻量约定式） ----------------

_DOCX_QUESTION_HEAD = re.compile(r"^\s*(\d+)[\.、)\s]\s*(?:【([^】]*)】)?\s*(.*)$")
_DOCX_OPTION_LINE = re.compile(r"^\s*([A-Za-z])[\.、)\s]\s*(.+)$")
_DOCX_ANSWER_LINE = re.compile(r"^\s*(?:答案|正确答案|参考答案)[:：]\s*(.+)$")
_DOCX_EXPLAIN_LINE = re.compile(r"^\s*(?:解析|解答|说明)[:：]\s*(.+)$")
# 同时识别「分值:5 / 分数 5」「5分」两种写法
_DOCX_SCORE_LABELED = re.compile(r"(?:分值|分数)[:：]?\s*(\d+(?:\.\d+)?)")
_DOCX_SCORE_BARE = re.compile(r"(\d+(?:\.\d+)?)\s*分")


def _split_marker(marker_inside: str) -> tuple[str, str]:
    """从 【...】 中提取 (题型, 分值)。

    支持：单选 / 单选,5分 / 单选 5分 / 单选,分值:5 / 单选；分数 5 等。
    """
    text = (marker_inside or "").strip()
    if not text:
        return "", ""
    score = ""
    m = _DOCX_SCORE_LABELED.search(text)
    if m:
        score = m.group(1)
        text = _DOCX_SCORE_LABELED.sub("", text)
    else:
        m = _DOCX_SCORE_BARE.search(text)
        if m:
            score = m.group(1)
            text = _DOCX_SCORE_BARE.sub("", text)
    qtype = text.strip(" ,，;；、\t")
    return qtype, score


def parse_docx(file_bytes: bytes) -> list[ParsedRow]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - 运行时给清晰报错
        raise RuntimeError("解析 Word 需要安装 python-docx：pip install python-docx") from exc

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [(p.text or "").rstrip() for p in doc.paragraphs]

    # 只在出现第一个「序号开头」之后才开始收集，丢掉前面的标题/填写说明等
    blocks: list[list[str]] = []
    current: list[str] = []
    seen_first_head = False
    for line in paragraphs:
        if not line.strip():
            continue
        is_head = bool(_DOCX_QUESTION_HEAD.match(line))
        if is_head:
            if current:
                blocks.append(current)
            current = [line]
            seen_first_head = True
        elif seen_first_head:
            current.append(line)
        # else: 第一题之前的内容，忽略
    if current:
        blocks.append(current)

    out: list[ParsedRow] = []
    for i, block in enumerate(blocks):
        head = _DOCX_QUESTION_HEAD.match(block[0])
        if not head:
            out.append(ParsedRow(idx=i, ok=False, errors=["题块未以序号开头"], raw={"lines": block}))
            continue
        marker_inside = head.group(2) or ""
        stem_first = head.group(3) or ""

        qtype_raw, score = _split_marker(marker_inside)

        options: list[str] = []
        stem_lines: list[str] = [stem_first] if stem_first else []
        correct_raw = ""
        explanation = ""
        for line in block[1:]:
            opt_m = _DOCX_OPTION_LINE.match(line)
            ans_m = _DOCX_ANSWER_LINE.match(line)
            exp_m = _DOCX_EXPLAIN_LINE.match(line)
            if opt_m:
                options.append(opt_m.group(2).strip())
            elif ans_m:
                correct_raw = ans_m.group(1).strip()
            elif exp_m:
                explanation = exp_m.group(1).strip()
            else:
                stem_lines.append(line.strip())

        if not qtype_raw:
            # 启发式推断
            if options and not correct_raw:
                qtype_raw = "single"
            elif not options and correct_raw in {"对", "错", "是", "否", "T", "F"}:
                qtype_raw = "judge"
            elif options and correct_raw and len(re.split(r"[\s,，、]+", correct_raw)) > 1:
                qtype_raw = "multiple"
            elif not options:
                qtype_raw = "short_answer"
            else:
                qtype_raw = "single"

        raw = {
            "question_type": qtype_raw,
            "stem": "\n".join(s for s in stem_lines if s),
            "options": options,
            "correct_answer": correct_raw,
            "score": score or 5,
            "category": "",
            "tag": "",
            "difficulty": "",
            "explanation": explanation,
        }
        ok, errors, normalized = validate_row(raw)
        out.append(ParsedRow(idx=i, ok=ok, data=normalized, errors=errors, raw=raw))
    return out


# ---------------- Excel 模板生成 ----------------


def build_excel_template() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "题库导入模板"

    header_fill = PatternFill(start_color="FFE6F0FA", end_color="FFE6F0FA", fill_type="solid")
    required_fill = PatternFill(start_color="FFFFF1F0", end_color="FFFFF1F0", fill_type="solid")
    bold = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for col_idx, header in enumerate(_EXCEL_HEADERS, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = bold
        cell.alignment = center
        cell.fill = required_fill if header.endswith("*") else header_fill
        ws.column_dimensions[get_column_letter(col_idx)].width = 18 if header == "题干*" else 14

    ws.column_dimensions["B"].width = 38

    samples = [
        ["单选", "下列哪一项是销售工作的核心？", "产品", "客户", "渠道", "价格", "", "", "B", 5, "销售基础", "客户", "简单", "以客户为中心是销售第一原则"],
        ["多选", "下列属于商务礼仪的是？", "守时", "得体着装", "随意打断", "称谓恰当", "", "", "AB,D", 5, "商务礼仪", "礼仪", "中等", "随意打断不是好的礼仪"],
        ["判断", "客户异议必须立即驳斥。", "对", "错", "", "", "", "", "错", 3, "客户管理", "异议处理", "简单", "应先理解再回应"],
        ["填空", "SPIN 销售法的四类问题是 ____、____、暗示问题、需求-效益问题。", "", "", "", "", "", "", "情境问题|难点问题", 4, "销售方法", "SPIN", "中等", "S-P-I-N 四类问题"],
        ["简答", "请描述一次成功的客户拜访前你会做哪些准备？", "", "", "", "", "", "", "", 10, "客户拜访", "准备", "中等", "可参考：客户背景、需求、产品方案、应对预案"],
    ]
    for r, row in enumerate(samples, start=2):
        for c, value in enumerate(row, start=1):
            ws.cell(row=r, column=c, value=value)

    notes = [
        "",
        "填写说明：",
        "1) 题型支持：单选 / 多选 / 判断 / 填空 / 简答（也可写英文 single/multiple/judge/blank/short）",
        "2) 多选正确答案：AB 或 A,B 或 A、B 都可；单选只填一个字母",
        "3) 判断题：答案填 对/错（也接受 是/否、T/F）",
        "4) 填空题：多个空用 | 分隔，例如 北京|上海",
        "5) 简答题：可不填正确答案；提交时进入人工评分",
        "6) 分值：必须为正数；分类/标签/难度/解析为可选项",
    ]
    for i, line in enumerate(notes, start=len(samples) + 3):
        cell = ws.cell(row=i, column=1, value=line)
        ws.merge_cells(start_row=i, end_row=i, start_column=1, end_column=14)
        if line.startswith("填写说明"):
            cell.font = Font(bold=True, color="FF1677FF")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------- Word 模板生成 ----------------


def build_docx_template() -> bytes:
    """生成 Word (.docx) 题库导入模板：包含填写说明 + 5 类示例题。

    约定格式（与 parse_docx 配套）：
      - 每题以「序号. 内容」开头（也接受 1、 / 1) / 1<空格>）
      - 在序号后用【题型,分值N分】标注（可省略，省略时由启发式推断）
      - 选项用「A. 内容」「B. 内容」… 单独成行
      - 「答案：X」标注正确答案；多选写 AB 或 A,B；判断写 对/错；填空多空用 | 分隔
      - 「解析：…」可选
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("生成 Word 模板需要安装 python-docx：pip install python-docx") from exc

    doc = Document()
    # 标题
    title = doc.add_heading("题库批量导入 · Word 模板", level=1)
    for run in title.runs:
        run.font.size = Pt(18)

    # 填写说明
    doc.add_heading("填写说明", level=2)
    notes = [
        "1) 每题以「序号. 题干」开头，紧跟在序号后可写【题型,分值N分】，例如：1. 【单选,5分】下列哪一项……",
        "2) 题型支持：单选 / 多选 / 判断 / 填空 / 简答（也接受 single/multiple/judge/blank/short）",
        "3) 选项另起一行，格式 A. 选项内容 / B. 选项内容（最多 6 个，A–F）",
        "4) 用「答案：…」标注正确答案；多选写 AB 或 A,B；判断写 对/错；填空多空用 | 分隔",
        "5) 解析可选，使用「解析：…」标注",
        "6) 题与题之间空一行，方便阅读；上传时空行会被忽略",
    ]
    for line in notes:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(10.5)

    # 示例题
    doc.add_heading("示例题（可直接修改）", level=2)
    samples = [
        [
            "1. 【单选,5分】下列哪一项是销售工作的核心？",
            "A. 产品",
            "B. 客户",
            "C. 渠道",
            "D. 价格",
            "答案：B",
            "解析：以客户为中心是销售第一原则。",
        ],
        [
            "2. 【多选,5分】下列属于商务礼仪的是？",
            "A. 守时",
            "B. 得体着装",
            "C. 随意打断",
            "D. 称谓恰当",
            "答案：A,B,D",
            "解析：随意打断不属于商务礼仪。",
        ],
        [
            "3. 【判断,3分】客户异议必须立即驳斥。",
            "答案：错",
            "解析：应先理解再回应。",
        ],
        [
            "4. 【填空,4分】SPIN 销售法的四类问题是 ____、____、暗示问题、需求-效益问题。",
            "答案：情境问题|难点问题",
            "解析：S-P-I-N 四类问题。",
        ],
        [
            "5. 【简答,10分】请描述一次成功的客户拜访前你会做哪些准备？",
            "答案：可参考 客户背景 / 需求 / 产品方案 / 应对预案 等维度展开。",
        ],
    ]
    for lines in samples:
        for line in lines:
            doc.add_paragraph(line)
        doc.add_paragraph("")  # 题间空行

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------- 序列化辅助 ----------------


def rows_to_json(rows: list[ParsedRow]) -> str:
    return json.dumps([r.to_dict() for r in rows], ensure_ascii=False)


def rows_from_json(text: str) -> list[ParsedRow]:
    data = json.loads(text or "[]")
    return [
        ParsedRow(
            idx=int(item.get("idx", i)),
            ok=bool(item.get("ok")),
            data=item.get("data"),
            errors=list(item.get("errors", []) or []),
            raw=item.get("raw"),
        )
        for i, item in enumerate(data)
    ]
